# Copyright (C) 2026 Intel Corporation
# SPDX-License-Identifier: Apache-2.0

import asyncio
import os
import re
import unicodedata
from bs4 import BeautifulSoup
from typing import List, Dict, Any

from comps import CustomLogger
from comps.dataprep.config import get_dataprep_llm_config, get_llm_extra_body, get_client
from comps.dataprep.utils import extract_json_from_response

logger = CustomLogger("dataprep-doc-hierarchy-parser", os.getenv("LOG_LEVEL", "INFO"))


class DocHierarchyParser:
    """
    国际化文档层级解析器 (支持 cn, en, th)
    自动识别标题层级、过滤目录、构建树状结构
    """

    DEFAULT_CONTEXT_TOKENS = int(os.getenv("DATAPREP_DOC_PARSER_CONTEXT_TOKENS", "32768"))
    DEFAULT_PROMPT_RESERVED_TOKENS = int(os.getenv("DATAPREP_DOC_PARSER_RESERVED_TOKENS", "8192"))
    DEFAULT_MAX_INPUT_TOKENS = int(os.getenv("DATAPREP_DOC_PARSER_MAX_INPUT_TOKENS", "4096"))
    DEFAULT_CHUNK_OVERLAP_TOKENS = int(os.getenv("DATAPREP_DOC_PARSER_OVERLAP_TOKENS", "1024"))
    DEFAULT_MAX_CHUNKS = int(os.getenv("DATAPREP_DOC_PARSER_MAX_CHUNKS", "32"))
    DEFAULT_CONTEXT_RETRY_DEPTH = int(os.getenv("DATAPREP_DOC_PARSER_CONTEXT_RETRY_DEPTH", "3"))

    # # 语言配置规则库
    # LANG_CONFIG = {
    #     "cn": {
    #         # 匹配：1.1 标题, 一、标题, 第一章 标题
    #         "heading": r"^((?:\d+(?:\.\d+)*\.?)|第[一二三四五六七八九十百]+[章节]|(?<!\d)[一二三四五六七八九十百]+)[.、\s]+([^\n\d\)\）]{2,}.*)$",
    #         "toc_indicator": r"[\.·…\-\s–—]{3,}\d+\s*$"
    #     },
    #     "en": {
    #         # 匹配：Chapter 1, Section I, 1.1 Heading (需大写字母开头避免误判正文)
    #         "heading": r"^((?:(?:Chapter|Section|Part|Article)\s+)?(?:\d+|[IVXLCDM]+)(?:\.\d+)*\.?)[.\s]+([A-Z][^\n]{2,})$",
    #         "toc_indicator": r"[\.·…\-\s–—]{3,}\d+\s*$"
    #     },
    #     "th": {
    #         # 匹配：บทที่ ๑ (第1章), ๑.๑ 标题, 1.1 标题
    #         "heading": r"^((?:บทที่|ส่วนที่|ข้อ)?\s*(?:[\d|๑-๙]+(?:[\.][\d|๑-๙]+)*\.?))[.\s]+([^\n]+)$",
    #         "toc_indicator": r"[\.·…\-\s–—]{3,}[\d|๑-๙]+\s*$"
    #     }
    # }

    # 语言配置规则库
    LANG_CONFIG = {
        "cn": {
            # 匹配：1.1 标题, 一、标题, 第一章 标题
            "heading": r"^(?!.*\|)(?:\d+\s|第[一二三四五六七八九十百]+[章节]|[一二三四五六七八九十百]+、)(.+)$",
            "toc_indicator": r"[\.·…\-\s–—]{3,}\d+\s*$"
        },
        "en": {
            # 匹配：Chapter 1, Section I, 1.1 Heading (需大写字母开头避免误判正文)
            "heading" : r"^(?!.*\|)(?:(?:Chapter\s+(?:\d+|[IVXLCDM]+))|(?:Section\s+[IVXLCDM]+)|\d+\.\s|\d+\s)(.+)$",
            "toc_indicator": r"[\.·…\-\s–—]{3,}\d+\s*$"
        },
        "th": {
            # 匹配：บทที่ ๑ (第1章), ๑.๑ 标题, 1.1 标题
            "heading": r"^(?!.*\|)(?:\d+\.\s|\d+\s|[\u0E50-\u0E59]+、|[\u0E00-\u0E7F]+、)(.+)$",
            "toc_indicator": r"[\.·…\-\s–—]{3,}[\d|๑-๙]+\s*$"
        }
    }

    def __init__(self, lang: str = "cn", ignore_case: bool = True):
        self.lang = lang.lower()
        config = self.LANG_CONFIG.get(self.lang, self.LANG_CONFIG["cn"])

        flags = re.M | (re.I if ignore_case else 0)
        self.pattern = re.compile(config["heading"], flags)
        # self.toc_pattern = re.compile(config["toc_indicator"], flags)

    def parse(self, text: str) -> List[Dict[str, Any]]:
        """解析文本返回树状结构"""
        sections = self._extract_sections(text)
        if not sections:
            return [{"total_title": "Full Text", "title_content": "Full Text","content":text.strip()}]
        return sections

    async def aparse(
        self,
        text: str,
        client=None,
        use_llm: bool = True,
    ) -> List[Dict[str, Any]]:
        """
        异步解析入口:
        - 默认优先使用 LLM 做一级标题抽取
        - 若模型解析失败/无结果，则回退到正则解析
        """
        if not use_llm:
            return self.parse(text)

        try:
            sections = await self._extract_sections_by_llm(
                text=text,
                client=client,
            )
            if sections:
                return sections
        except Exception as exc:
            logger.warning(f"LLM 文档层级解析失败，回退正则方案: {exc}")

        return self.parse(text)

    def _extract_sections(self, text: str) -> List[Dict[str, Any]]:
        """提取所有标题及其正文片段"""
        matches = list(self.pattern.finditer(text))
        sections = []

        for i, match in enumerate(matches):
            # 过滤包含表格线的内容
            if "|" in match.group(0):
                continue

            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            total_title = match.group(0).strip()
            title_content = match.group(1).strip()

            sections.append({
                "total_title": total_title,
                'title_content':title_content,
                "content": content,
            })
        return sections

    async def _extract_sections_by_llm(
        self,
        text: str,
        client=None,
    ) -> List[Dict[str, Any]]:
        if not text or not text.strip():
            return []

        llm_client = client or get_client()
        resolved_model_name = self._resolve_model_name()
        chunks = self._split_text_for_llm(text)
        if not chunks:
            return []

        logger.info("开始 LLM 一级标题列表提取")
        collected_titles: list[dict[str, str]] = []
        for idx, chunk in enumerate(chunks, start=1):
            chunk_titles = await self._extract_titles_for_chunk(
                client=llm_client,
                model_name=resolved_model_name,
                chunk=chunk,
                chunk_index=idx,
                total_chunks=len(chunks),
            )
            collected_titles.extend(chunk_titles)

        merged_titles = self._merge_titles(collected_titles)
        sections = self._split_sections_by_titles(text, merged_titles)
        if sections:
            return sections
        return []

    def _split_text_for_llm(self, text: str) -> List[str]:
        max_input_tokens = self.DEFAULT_CONTEXT_TOKENS - self.DEFAULT_PROMPT_RESERVED_TOKENS
        if max_input_tokens <= 0:
            max_input_tokens = 4096
        if self.DEFAULT_MAX_INPUT_TOKENS > 0:
            max_input_tokens = min(max_input_tokens, self.DEFAULT_MAX_INPUT_TOKENS)
        lines = text.splitlines()
        if not lines:
            return [text.strip()] if text.strip() else []

        chunks: list[str] = []
        current_lines: list[str] = []
        current_tokens = 0

        for line in lines:
            estimated = self._estimate_tokens(line) + 1

            if current_lines and current_tokens + estimated > max_input_tokens:
                chunks.append("\n".join(current_lines).strip())
                overlap_lines = self._collect_overlap_lines(
                    current_lines,
                    self.DEFAULT_CHUNK_OVERLAP_TOKENS,
                )
                current_lines = overlap_lines[:]
                current_tokens = sum(self._estimate_tokens(item) + 1 for item in current_lines)

            if estimated > max_input_tokens:
                for part in self._split_long_line(line, max_input_tokens):
                    if current_lines and current_tokens + self._estimate_tokens(part) + 1 > max_input_tokens:
                        chunks.append("\n".join(current_lines).strip())
                        overlap_lines = self._collect_overlap_lines(
                            current_lines,
                            self.DEFAULT_CHUNK_OVERLAP_TOKENS,
                        )
                        current_lines = overlap_lines[:]
                        current_tokens = sum(self._estimate_tokens(item) + 1 for item in current_lines)
                    current_lines.append(part)
                    current_tokens += self._estimate_tokens(part) + 1
                continue

            current_lines.append(line)
            current_tokens += estimated

        if current_lines:
            chunks.append("\n".join(current_lines).strip())

        cleaned_chunks = [chunk for chunk in chunks if chunk]
        if len(cleaned_chunks) > self.DEFAULT_MAX_CHUNKS:
            logger.warning(
                f"文档切片数量过多，已截断为前 {self.DEFAULT_MAX_CHUNKS} 片，原始片数={len(cleaned_chunks)}"
            )
            cleaned_chunks = cleaned_chunks[:self.DEFAULT_MAX_CHUNKS]
        return cleaned_chunks

    def _collect_overlap_lines(self, lines: List[str], overlap_budget: int) -> List[str]:
        overlap_lines: list[str] = []
        tokens = 0
        for line in reversed(lines):
            tokens += self._estimate_tokens(line) + 1
            overlap_lines.insert(0, line)
            if tokens >= overlap_budget:
                break
        return overlap_lines

    def _split_long_line(self, line: str, max_input_tokens: int) -> List[str]:
        parts: list[str] = []
        cursor = 0
        while cursor < len(line):
            end = min(len(line), cursor + max_input_tokens)
            parts.append(line[cursor:end])
            cursor = end
        return parts

    def _estimate_tokens(self, text: str) -> int:
        if not text:
            return 0
        units = re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9_]+|[^\s]", text)
        return len(units)

    async def _extract_titles_for_chunk(
        self,
        client,
        model_name: str,
        chunk: str,
        chunk_index: int,
        total_chunks: int,
        split_depth: int = 0,
    ) -> List[Dict[str, str]]:
        prompt = self._build_title_prompt(
            chunk=chunk,
            chunk_index=chunk_index,
            total_chunks=total_chunks,
        )

        llm_config = get_dataprep_llm_config()
        fallback_model_name = llm_config.model if llm_config.model and llm_config.model != model_name else ""
        try:
            response = await self._create_chunk_completion(
                client=client,
                model_name=model_name,
                prompt=prompt,
            )
        except Exception as exc:
            if fallback_model_name and self._is_model_not_found_error(exc):
                logger.warning(
                    f"文档切片解析模型不存在，自动回退到当前 dataprep 模型: from={model_name}, to={fallback_model_name}, chunk={chunk_index}/{total_chunks}"
                )
                response = await self._create_chunk_completion(
                    client=client,
                    model_name=fallback_model_name,
                    prompt=prompt,
                )
            else:
                if self._is_context_length_error(exc):
                    if split_depth >= self.DEFAULT_CONTEXT_RETRY_DEPTH:
                        logger.error(
                            f"LLM 文档切片解析失败(超出上下文且达到重试上限): chunk={chunk_index}/{total_chunks}, model={model_name}, depth={split_depth}, error={exc}"
                        )
                        raise

                    sub_chunks = self._split_chunk_for_context_retry(chunk)
                    if len(sub_chunks) <= 1:
                        logger.error(
                            f"LLM 文档切片解析失败(超出上下文且无法继续切分): chunk={chunk_index}/{total_chunks}, model={model_name}, error={exc}"
                        )
                        raise

                    logger.warning(
                        f"LLM 文档切片超出模型上下文，自动切分重试: chunk={chunk_index}/{total_chunks}, model={model_name}, depth={split_depth + 1}, sub_chunks={len(sub_chunks)}"
                    )
                    merged_titles: list[dict[str, str]] = []
                    for sub_idx, sub_chunk in enumerate(sub_chunks, start=1):
                        sub_titles = await self._extract_titles_for_chunk(
                            client=client,
                            model_name=model_name,
                            chunk=sub_chunk,
                            chunk_index=sub_idx,
                            total_chunks=len(sub_chunks),
                            split_depth=split_depth + 1,
                        )
                        merged_titles.extend(sub_titles)
                    return merged_titles

                logger.error(
                    f"LLM 文档切片解析失败: chunk={chunk_index}/{total_chunks}, model={model_name}, error={exc}"
                )
                raise

        parsed = extract_json_from_response(response.choices[0].message.content)
        return self._sanitize_titles_response(parsed)

    async def _create_chunk_completion(self, client, model_name: str, prompt: str):
        extra_body = get_llm_extra_body(model_name)
        return await asyncio.to_thread(
            client.chat.completions.create,
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0,
            extra_body=extra_body
        )

    def _resolve_model_name(self) -> str:
        llm_config = get_dataprep_llm_config()
        return llm_config.model

    def _is_model_not_found_error(self, exc: Exception) -> bool:
        message = str(exc).lower()
        return "does not exist" in message or "notfounderror" in message or "404" in message

    def _is_context_length_error(self, exc: Exception) -> bool:
        message = str(exc).lower()
        return (
            "maximum context length" in message
            or "reduce the length of the messages" in message
            or "context length is" in message
            or "requested" in message and "tokens" in message
        )

    def _split_chunk_for_context_retry(self, chunk: str) -> List[str]:
        text = chunk.strip()
        if not text:
            return []

        lines = text.splitlines()
        if len(lines) >= 2:
            mid = len(lines) // 2
            left = "\n".join(lines[:mid]).strip()
            right = "\n".join(lines[mid:]).strip()
            return [part for part in [left, right] if part]

        mid = len(text) // 2
        left = text[:mid].strip()
        right = text[mid:].strip()
        return [part for part in [left, right] if part]

    def _build_title_prompt(
        self,
        chunk: str,
        chunk_index: int,
        total_chunks: int,
    ) -> str:
        language_name = {
            "cn": "中文",
            "zh": "中文",
            "en": "英文",
            "th": "泰文",
        }.get(self.lang, self.lang)
        return (
            "你是文档结构助手，只提取“一级标题列表”，不要输出正文。\n"
            "要求：\n"
            "1. 仅保留一级标题，不要输出二级及以下标题。\n"
            "2. total_title 必须尽量保持原文标题行（包含编号）。\n"
            "3. title_content 为去掉编号后的标题文本。\n"
            "4. 按在当前切片中出现顺序输出。\n"
            "5. 不要输出解释，只输出 JSON。\n\n"
            "输出格式：\n"
            "{\n"
            '  "titles": [\n'
            '    {"total_title": "1. 目的", "title_content": "目的"}\n'
            "  ]\n"
            "}\n\n"
            f"文档语言：{language_name}\n"
            f"当前切片：第 {chunk_index} / {total_chunks} 片\n\n"
            "待解析文本：\n"
            f"{chunk}"
        )

    def _sanitize_titles_response(self, parsed: Any) -> List[Dict[str, str]]:
        if isinstance(parsed, dict):
            raw_titles = parsed.get("titles", [])
        elif isinstance(parsed, list):
            raw_titles = parsed
        else:
            raw_titles = []

        sanitized_titles: list[dict[str, str]] = []
        for item in raw_titles or []:
            if not isinstance(item, dict):
                continue
            total_title = str(item.get("total_title", "") or "").strip()
            title_content = str(item.get("title_content", "") or "").strip()
            if not total_title and title_content:
                total_title = title_content
            if not total_title:
                continue
            if not title_content:
                title_content = self._strip_numbering_prefix(total_title)

            sanitized_titles.append({
                "total_title": total_title,
                "title_content": title_content or total_title,
            })

        return sanitized_titles

    def _merge_titles(self, titles: List[Dict[str, str]]) -> List[Dict[str, str]]:
        merged: list[dict[str, str]] = []
        for item in titles:
            total_title = item.get("total_title", "").strip()
            title_content = item.get("title_content", "").strip() or self._strip_numbering_prefix(total_title)
            if not total_title:
                continue

            key = self._normalize_for_match(total_title)
            if not key:
                continue
            if merged and self._normalize_for_match(merged[-1]["total_title"]) == key:
                continue
            merged.append({
                "total_title": total_title,
                "title_content": title_content or total_title,
            })
        return merged

    def _split_sections_by_titles(self, text: str, titles: List[Dict[str, str]]) -> List[Dict[str, Any]]:
        if not text.strip():
            return []
        if not titles:
            return []

        title_hits = self._locate_titles_in_text(text, titles)
        if not title_hits:
            return []

        sections: list[dict[str, Any]] = []
        preamble = self._strip_toc_lines(text[:title_hits[0]["title_start"]]).strip()

        for i, hit in enumerate(title_hits):
            next_start = title_hits[i + 1]["title_start"] if i + 1 < len(title_hits) else len(text)
            content = text[hit["title_end"]:next_start].strip()
            if i == 0 and preamble:
                content = f"{preamble}\n{content}".strip()
            sections.append({
                "total_title": hit["total_title"],
                "title_content": hit["title_content"],
                "content": content,
            })

        return sections

    def _locate_titles_in_text(self, text: str, titles: List[Dict[str, str]]) -> List[Dict[str, Any]]:
        lines = []
        offset = 0
        toc_mode = False

        for raw_line in text.splitlines(keepends=True):
            line = raw_line.rstrip("\r\n")
            start = offset
            end = offset + len(line)
            offset += len(raw_line)
            line_is_toc_title = self._is_toc_title_line(line)
            line_is_toc_entry = self._is_toc_entry_line(line)
            is_toc_line = False

            if line_is_toc_title:
                toc_mode = True
                is_toc_line = True
            elif toc_mode:
                if not line.strip():
                    is_toc_line = True
                elif line_is_toc_entry:
                    is_toc_line = True
                else:
                    toc_mode = False
            else:
                # 无显式“目录”标题时，只在文档前部把目录项当作候选过滤
                if len(lines) < 120 and line_is_toc_entry:
                    is_toc_line = True

            lines.append({
                "line": line,
                "start": start,
                "end": end,
                "normalized": self._normalize_for_match(line),
                "is_toc": is_toc_line,
            })

        hits: list[dict[str, Any]] = []
        cursor = 0

        for title in titles:
            candidates = self._build_title_candidates(
                total_title=title.get("total_title", ""),
                title_content=title.get("title_content", ""),
            )
            best_idx = -1
            best_score = -1

            for idx in range(cursor, len(lines)):
                if lines[idx]["is_toc"]:
                    continue

                line_norm = lines[idx]["normalized"]
                if not line_norm:
                    continue

                score_for_line = -1
                for cand in candidates:
                    if not cand:
                        continue
                    if line_norm == cand:
                        score_for_line = max(score_for_line, 3)
                    elif line_norm.startswith(cand):
                        score_for_line = max(score_for_line, 2)
                    elif len(cand) >= 4 and cand in line_norm:
                        score_for_line = max(score_for_line, 1)

                if score_for_line > best_score:
                    best_score = score_for_line
                    best_idx = idx
                    if best_score == 3:
                        break

            if best_idx == -1:
                continue

            line = lines[best_idx]
            hits.append({
                "total_title": title.get("total_title", "").strip(),
                "title_content": title.get("title_content", "").strip() or self._strip_numbering_prefix(title.get("total_title", "").strip()),
                "title_start": line["start"],
                "title_end": line["end"],
            })
            cursor = best_idx + 1

        return hits

    def _strip_toc_lines(self, text: str) -> str:
        if not text:
            return text

        output_lines = []
        toc_mode = False
        for line in text.splitlines():
            if self._is_toc_title_line(line):
                toc_mode = True
                continue

            if toc_mode:
                if not line.strip() or self._is_toc_entry_line(line):
                    continue
                toc_mode = False

            if self._is_toc_entry_line(line):
                continue
            output_lines.append(line)

        return "\n".join(output_lines)

    def _is_toc_title_line(self, line: str) -> bool:
        normalized = unicodedata.normalize("NFKC", line or "").strip()
        if not normalized:
            return False

        toc_titles = {
            "cn": {"目录"},
            "zh": {"目录"},
            "en": {"contents", "table of contents"},
            "th": {"สารบัญ"},
        }
        title_set = toc_titles.get(self.lang, toc_titles["cn"])
        return normalized.lower() in title_set or normalized in title_set

    def _is_toc_entry_line(self, line: str) -> bool:
        normalized = unicodedata.normalize("NFKC", line or "").strip()
        if not normalized:
            return False

        # 典型目录点线 + 页码
        if re.search(r"(?:\.|·|…|-|—|–){3,}\s*\d{1,4}\s*$", normalized):
            return True

        # 目录中常见：编号标题 + 末尾页码
        if len(normalized) <= 120 and re.search(r"\s+\d{1,4}\s*$", normalized):
            if re.match(
                r"^(?:\d+(?:\.\d+)*[\.、]?\s*|第[一二三四五六七八九十百千]+[章节篇]?\s*|"
                r"[一二三四五六七八九十百千]+、\s*|(?:chapter|section|part|article)\s+(?:\d+|[ivxlcdm]+)\s*)",
                normalized,
                flags=re.I,
            ):
                return True

        return False

    def _build_title_candidates(self, total_title: str, title_content: str) -> List[str]:
        values = []
        for value in [total_title, title_content, self._strip_numbering_prefix(total_title)]:
            normalized = self._normalize_for_match(value)
            if normalized and normalized not in values:
                values.append(normalized)
        return values

    def _normalize_for_match(self, text: str) -> str:
        normalized = unicodedata.normalize("NFKC", text or "")
        normalized = normalized.replace("\u3000", " ")
        normalized = re.sub(r"\s+", "", normalized)
        return normalized.lower()

    def _strip_numbering_prefix(self, text: str) -> str:
        if not text:
            return ""
        cleaned = text.strip()
        cleaned = re.sub(
            r"^(?:\d+(?:\.\d+)*[\.、]?\s*|第[一二三四五六七八九十百千]+[章节篇]\s*|[一二三四五六七八九十百千]+、\s*)",
            "",
            cleaned,
        )
        cleaned = re.sub(
            r"^(?:chapter|section|part|article)\s+(?:\d+|[ivxlcdm]+)[\.\s]*",
            "",
            cleaned,
            flags=re.I,
        )
        return cleaned.strip()

    def _normalize_title(self, title: str) -> str:
        return re.sub(r"\s+", "", title or "").lower()

    def _merge_content(self, existing: str, new_text: str) -> str:
        existing = (existing or "").strip()
        new_text = (new_text or "").strip()

        if not existing:
            return new_text
        if not new_text:
            return existing
        if new_text in existing:
            return existing

        existing_lines = [line.strip() for line in existing.splitlines() if line.strip()]
        new_lines = [line.strip() for line in new_text.splitlines() if line.strip()]

        max_overlap = min(len(existing_lines), len(new_lines), 20)
        overlap_size = 0
        for size in range(max_overlap, 0, -1):
            if existing_lines[-size:] == new_lines[:size]:
                overlap_size = size
                break

        if overlap_size:
            merged_lines = existing_lines + new_lines[overlap_size:]
            return "\n".join(merged_lines).strip()

        if existing.endswith(new_text):
            return existing

        return f"{existing}\n{new_text}".strip()

    def _calculate_level(self, num: str) -> int:
        """根据编号计算层级深度"""
        # 1. 优先处理多级数字点号：1.1 (Level 2), 1.1.1 (Level 3)
        dot_count = num.count(".")
        if dot_count > 0:
            # 如果是 "1." 结尾则算 1 级，如果是 "1.1" 则算 2 级
            return dot_count + (0 if num.endswith(".") else 1)

        # 2. 识别各语言的一级标题前缀
        lvl1_prefixes = {
            "cn": ["第", "一", "二", "三"],
            "en": ["Chapter", "Part", "Section", "Article", "I", "II", "III", "IV", "V"],
            "th": ["บทที่", "ส่วนที่", "ข้อ", "๑", "๒", "๓"]
        }

        for p in lvl1_prefixes.get(self.lang, []):
            if num.startswith(p):
                return 1
        return 1

    def _build_hierarchy(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """线性列表转树状结构"""
        root = []
        stack = []
        for sec in sections:
            while stack and stack[-1]["level"] >= sec["level"]:
                stack.pop()
            if not stack:
                root.append(sec)
            else:
                stack[-1]["children"].append(sec)
            stack.append(sec)
        return root


class MarkdownHierarchyParser(DocHierarchyParser):

    def __init__(self, lang: str = "cn", ignore_case: bool = True):
        """
        :param ignore_case: 是否忽略大小写（默认True）
        """
        # 1. 调用父类初始化，获取基础语言配置
        super().__init__(lang, ignore_case)

        # 2. 增强正则表达式：允许开头有可选的 Markdown 井号 (?:#{1,6}\s*)?
        config = self.LANG_CONFIG.get(self.lang, self.LANG_CONFIG["cn"])
        original_heading_regex = config["heading"].lstrip("^")

        # 组合后的正则：支持 "# 1.1 标题" 或 "1.1 标题"
        self.md_pattern = re.compile(
            r"^(#{1,6}\s+)??" + original_heading_regex,
            re.M | (re.I if ignore_case else 0)
        )

        # 3. 多语言目录标题识别（用于过滤掉整个目录块）
        # 使用更宽松的匹配，[\s\u3000\u00a0]* 匹配各种空格变体（半角、全角、不间断空格）
        self.catalog_titles = {
            "cn": r"^[\s\u3000\u00a0]*#*[\s\u3000\u00a0]*目[\s\u3000\u00a0]*录[\s\u3000\u00a0]*$",
            "en": r"^[\s\u3000\u00a0]*#*[\s\u3000\u00a0]*(Contents|Table[\s\u3000\u00a0]*of[\s\u3000\u00a0]*Contents)[\s\u3000\u00a0]*$",
            "th": r"^[\s\u3000\u00a0]*#*[\s\u3000\u00a0]*สารบัญ[\s\u3000\u00a0]*$"
        }
        self.cat_regex = re.compile(
            self.catalog_titles.get(self.lang, self.catalog_titles["cn"]),
            re.I | re.M
        )

    def parse(self, text: str) -> List[Dict[str, Any]]:
        """
        解析逻辑：先过滤目录块，再提取层级
        """
        # 第一步：过滤目录块（如文档开头手动写的目录列表）
        filtered_text = self._filter_catalog_block(text)

        # 第二步：提取章节
        sections = self._extract_md_sections(filtered_text)

        if not sections:
            return [{"title": "Full Text", "content": filtered_text.strip(), "level": 0, "children": []}]

        # 第三步：构建树状结构
        return self._build_hierarchy(sections)

    def _filter_catalog_block(self, text: str) -> str:
        """
        识别并删除目录章节块
        
        策略：
        1. 明确的目录标题（如 # 目录）后进入跳过模式
        2. 跳过模式下，只有遇到带 # 前缀的正文标题才退出
        3. 即使没有明确目录标题，也尝试检测连续的目录项特征
        """
        lines = text.splitlines()
        filtered_lines = []
        skip_mode = False
        
        # 目录项特征正则：以数字开头的标题行（无 # 前缀）
        toc_item_pattern = re.compile(
            r'^[\s\u3000\u00a0]*\d+[\s\u3000\u00a0\.、]+[\u4e00-\u9fff]',  # 数字 + 空格/点 + 中文
            re.M
        )

        for i, line in enumerate(lines):
            line_strip = line.strip()
            
            # 如果匹配到目录标题，进入跳过模式
            if self.cat_regex.match(line):
                skip_mode = True
                continue

            if skip_mode:
                # 退出条件：遇到带 # 前缀的标题（正文开始）
                if re.match(r"^#{1,6}\s+", line):
                    skip_mode = False
                    # 不 continue，让这行被添加到 filtered_lines
                elif line_strip == "":
                    # 空行继续跳过
                    continue
                else:
                    # 其他内容（目录项）继续跳过
                    continue

            filtered_lines.append(line)

        return "\n".join(filtered_lines)

    def _is_toc_entry(self, title: str, content: str) -> bool:
        """
        检测是否为目录项（无明确"目录"标题的隐式目录）
        
        目录项特征：
        1. title 中包含子章节编号（如 2.1, 3.1.2 等）
        2. title 中包含页码指示符（连续点号+数字）
        3. title 中多个章节编号连在一起（目录行被拼接）
        4. content 为空
        """
        if content.strip():
            # 有正文内容的不是目录项
            return False
        
        # 检测1：title 中包含子章节编号模式 (如 .2.1, .3.2 等)
        # 这种模式表示目录中子标题被拼接到了主标题后面
        if re.search(r'\.\d+\.\d+', title):
            return True
        
        # 检测2：title 中包含页码指示（连续的点号/省略号 + 数字）
        if re.search(r'[\.·…]{2,}\s*\d+', title):
            return True
        
        # 检测3：title 中有多个独立的章节编号紧密排列（如 "2.1 工艺原理2.2 流程描述"）
        # 检测连续出现的 "数字.数字 中文" 模式
        chapter_pattern = r'\d+\.\d+\s*[\u4e00-\u9fff]+'
        matches = re.findall(chapter_pattern, title)
        if len(matches) >= 2:
            return True
        
        # 检测4：title 过长且包含多个中文短语（目录行特征）
        # 正常标题一般不会太长，目录行会把多个子项拼在一起
        if len(title) > 50:
            # 检查是否有多个数字编号
            num_pattern = r'\d+(?:\.\d+)*'
            nums = re.findall(num_pattern, title)
            if len(nums) >= 3:
                return True
        
        return False

    def _extract_md_sections(self, text: str) -> List[Dict[str, Any]]:
        """
        提取 Markdown 章节，支持多语言编号和 # 层级
        """
        matches = list(self.md_pattern.finditer(text))
        sections = []

        for i, match in enumerate(matches):
            # 排除表格行
            if "|" in match.group(0):
                continue

            # 获取正文范围
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()

            # match 结果组：
            # group(1): 可能存在的 # 部分
            # group(2): 编号部分 (1.1 / 第一章 / ๑.๑)
            # group(3): 标题文本部分
            hash_part = match.group(1)
            heading_num = match.group(2).strip()
            title = match.group(3).strip()

            # 过滤掉 TOC 指示行（带页码的行）
            if self.toc_pattern.search(title):
                continue

            # 过滤隐式目录项（无"目录"标题，但有目录特征）
            if self._is_toc_entry(title, content):
                continue

            # 计算层级：优先看 # 的数量，没有则看数字深度
            if hash_part and hash_part.strip().startswith("#"):
                level = hash_part.strip().count("#")
            else:
                level = self._calculate_level(heading_num)

            sections.append({
                "heading_num": heading_num,
                "title": title,
                "content": content,
                "level": level,
                "children": []
            })
        return sections

    def _calculate_level(self, num: str) -> int:
        """
        复用并增强父类的层级计算
        """
        # 泰文数字转换逻辑（如果需要更精确，可以把 ๑ 换成 1 再计算 dot_count）
        thai_num_map = str.maketrans("๑๒๓๔๕๖๗๘๙๐", "1234567890")
        normalized_num = num.translate(thai_num_map)

        # 调用父类逻辑处理点号分隔符
        return super()._calculate_level(normalized_num)


class WordTableExtractorLite:
    """
    提取 Word 文档中所有表格（含嵌套表格），带标题关联（基于加粗检测）。
    """

    def extract(self, file_path: str, doc) -> List[Dict[str, Any]]:
        elements = list(doc.element.body)  # 保留文档顺序（段落+表格混合）
        tables = []
        table_index = 0

        for i, elem in enumerate(elements):
            if elem.tag.endswith('tbl'):
                table_index += 1
                table = doc.tables[len(tables)]

                # 获取上方最近的“加粗段落标题”
                title = self._get_prev_bold_paragraph(elements, i, doc)

                table_info = self._extract_table(table, table_index)
                table_info["table_title"] = title or f"表格 {table_index}"
                tables.append(table_info)

        return tables

    def _get_prev_bold_paragraph(self, elements, table_idx: int, doc) -> str:
        """
        向上查找最近的“加粗”段落，作为表格标题。
        逻辑：
        - 优先返回加粗段落；
        - 若无加粗，则返回第一个非空文本；
        - 若到上一个表都未找到加粗标题，则以第一个非空段落为兜底。
        """
        from docx.oxml.text.paragraph import CT_P
        element_to_paragraph = {p._element: p for p in doc.paragraphs}

        first_fallback_text = None  # 记录第一个非空段落（兜底）

        for j in range(table_idx - 1, -1, -1):
            elem = elements[j]

            # 如果遇到上一个表格，停止搜索
            if elem.tag.endswith('tbl'):
                break

            if elem.tag.endswith('p') and elem in element_to_paragraph:
                p = element_to_paragraph[elem]
                text = p.text.strip() if p.text else ""
                if not text:
                    continue

                # 记录第一个非空文本（兜底用）
                if first_fallback_text is None:
                    first_fallback_text = text

                # 检查是否为加粗段落
                has_bold = any(
                    run.text.strip() and getattr(run, "bold", False)
                    for run in p.runs
                )

                if has_bold:
                    return text

        # 没有加粗标题则使用兜底段落
        return first_fallback_text or ""

    def _extract_table(self, table, index: int) -> Dict[str, Any]:
        cells = []
        seen_texts = set()

        for row in table.rows:
            for cell in row.cells:
                text_parts = []

                # 当前单元格普通段落
                for p in cell.paragraphs:
                    if p.text.strip():
                        text_parts.append(p.text.strip())

                # 嵌套表格递归提取
                for nested in cell.tables:
                    nested_text = self._table_to_text(nested)
                    if nested_text.strip():
                        text_parts.append(nested_text.strip())

                if text_parts:
                    full_text = "\n".join(text_parts).strip()
                    norm = full_text.replace("\n", "").replace(" ", "")
                    if norm and norm not in seen_texts:
                        seen_texts.add(norm)
                        cells.append(full_text)
                else:
                    cells.append("")

        return {
            "table_index": index,
            "cells": cells
        }

    def _table_to_text(self, table) -> str:
        """将嵌套表格内容转为文本"""
        rows_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            rows_text.append(row_text)
        return "\n".join(rows_text)


class MagicPdfTableExtractor:
    """
    从 magic-pdf 的 Markdown 输出中提取包含“应急 + 演练”或“应急 + 预案”的表格块。
    支持未加 # 的简短标题行识别。
    """

    def extract(self, md_text: str) -> List[Dict[str, Any]]:
        """
        提取 Markdown 中符合条件的表格段落。
        按标题块分段，标题到下一个标题之间的全部内容。
        """
        lines = md_text.splitlines()
        headings = []

        for i, line in enumerate(lines):
            line_strip = line.strip()

            # 标题识别（符合“应急+演练/预案”）
            if (("应急" in line_strip and "演练" in line_strip) or
                    ("应急" in line_strip and "预案" in line_strip) or
                    ("演练总结" in line_strip) or
                    ("应急物资检查表" in line_strip) or
                    ("存在问题整改记录" in line_strip)
            ):
                chinese_chars = re.findall(r"[\u4e00-\u9fff]", line_strip)
                if len(chinese_chars) <= 15 and not any(ch in line_strip for ch in [':', '：', '、']):
                    prev_line = lines[i - 1].strip() if i > 0 else ""
                    if prev_line == "":
                        headings.append((i, line_strip))

        if not headings:
            return []

        # 依标题范围提取完整内容块
        tables = []
        for idx, (line_idx, title) in enumerate(headings):
            start_line = line_idx
            end_line = headings[idx + 1][0] if idx + 1 < len(headings) else len(lines)
            content_block = "\n".join(lines[start_line:end_line]).strip()

            # 提取所有 HTML 表格（可能多个）
            html_blocks = re.findall(r"<html>.*?</html>", content_block, re.S)
            html_blocks = list(dict.fromkeys(html_blocks))  # 去重

            # 替换每个 HTML 为表格二维字符串
            for html in html_blocks:
                matrices = self.html_to_matrix(html)
                # 将二维数组转换为可阅读文本，例如：
                # ┌──────────┬────────────┐
                # │ 序号     │ 名称       │
                # │ 1        │ 对讲机     │
                # └──────────┴────────────┘
                table_str_list = []
                for table in matrices:
                    # 简单文本化二维表
                    table_lines = [" | ".join(row) for row in table]
                    table_str = "\n".join(table_lines)
                    table_str_list.append(table_str)
                table_text = "\n\n[表格内容]\n" + "\n\n---\n\n".join(table_str_list) + "\n\n[/表格内容]\n"

                # 替换原HTML
                content_block = content_block.replace(html, table_text)

            tables.append({
                "table_index": f"{idx + 1}",
                "table_title": title,
                "cells": [content_block],  # 替换后的完整内容
            })

        return tables

    @staticmethod
    def html_to_matrix(html_str: str) -> List[List[List[str]]]:
        """
        将 HTML 表格转为二维数组；支持多个 <table>。
        """
        soup = BeautifulSoup(html_str, "html.parser")
        matrices = []
        for t in soup.find_all("table"):
            rows = []
            for tr in t.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            if rows:
                matrices.append(rows)
        return matrices
